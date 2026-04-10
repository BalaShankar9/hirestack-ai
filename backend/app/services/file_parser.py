"""
File Parser Utility
Extracts text from PDF, DOCX, DOC, and TXT files with high accuracy.
Handles edge cases: multi-column layouts, tables, headers/footers, encoding issues.
"""
import io
import re
from typing import List, Optional

import pdfplumber
from docx import Document as DocxDocument


class FileParser:
    """Utility for extracting text from various file formats."""

    async def extract_text(
        self,
        file_contents: bytes,
        file_type: str,
    ) -> str:
        """Extract text from file contents based on file type."""
        file_type = file_type.lower().strip(".")

        if file_type == "pdf":
            return await self._extract_pdf(file_contents)
        elif file_type in ("docx", "doc"):
            return await self._extract_docx(file_contents)
        elif file_type == "txt":
            return await self._extract_txt(file_contents)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def _extract_pdf(self, file_contents: bytes) -> str:
        """Extract text from PDF file with layout-aware extraction."""
        text_parts: List[str] = []

        try:
            with pdfplumber.open(io.BytesIO(file_contents)) as pdf:
                for page in pdf.pages:
                    # Try regular text extraction first
                    page_text = page.extract_text(
                        x_tolerance=2,
                        y_tolerance=3,
                    )

                    # If regular extraction yields little text, try table extraction
                    if not page_text or len(page_text.strip()) < 20:
                        tables = page.extract_tables()
                        if tables:
                            table_texts = []
                            for table in tables:
                                for row in table:
                                    cells = [c.strip() for c in row if c and c.strip()]
                                    if cells:
                                        table_texts.append(" | ".join(cells))
                            if table_texts:
                                page_text = "\n".join(table_texts)

                    if page_text and page_text.strip():
                        # Clean up common PDF extraction artifacts
                        page_text = self._clean_pdf_text(page_text)
                        text_parts.append(page_text)

        except Exception as e:
            raise ValueError(
                f"Could not read PDF file: {str(e)}. "
                "The file may be corrupted, password-protected, or image-based (scanned)."
            )

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise ValueError(
                "Could not extract text from PDF. "
                "The file may be image-based (scanned). "
                "Please try uploading a text-based PDF, DOCX, or TXT file instead."
            )

        return full_text

    def _clean_pdf_text(self, text: str) -> str:
        """Clean common PDF extraction artifacts."""
        # Fix broken words from column extraction (e.g., "Soft- ware" → "Software")
        text = re.sub(r"(\w)-\s+(\w)", r"\1\2", text)
        # Fix ligatures
        text = text.replace("ﬁ", "fi").replace("ﬂ", "fl").replace("ﬀ", "ff")
        text = text.replace("ﬃ", "ffi").replace("ﬄ", "ffl")
        # Remove form feed characters
        text = text.replace("\f", "\n")
        # Normalize whitespace within lines
        lines = text.split("\n")
        cleaned = []
        for line in lines:
            line = re.sub(r"[ \t]+", " ", line).strip()
            cleaned.append(line)
        return "\n".join(cleaned)

    async def _extract_docx(self, file_contents: bytes) -> str:
        """Extract text from DOCX file including paragraphs, tables, headers."""
        try:
            doc = DocxDocument(io.BytesIO(file_contents))
        except Exception as e:
            raise ValueError(f"Could not open document file: {str(e)}")

        text_parts: List[str] = []

        # Extract header content (often contains name/contact)
        for section in doc.sections:
            header = section.header
            if header:
                for para in header.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())

        # Extract main body paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract from tables (common in resume formatting)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # Deduplicate cells (merged cells can repeat text)
                seen = set()
                unique_cells = []
                for c in cells:
                    if c not in seen:
                        seen.add(c)
                        unique_cells.append(c)
                if unique_cells:
                    text_parts.append(" | ".join(unique_cells))

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            raise ValueError(
                "Could not extract text from document. "
                "The file may be empty or contain only images."
            )

        return full_text

    async def _extract_txt(self, file_contents: bytes) -> str:
        """Extract text from TXT file with robust encoding detection."""
        # Try encodings in order of likelihood for resumes
        for encoding in ["utf-8", "utf-8-sig", "utf-16", "latin-1", "cp1252", "ascii"]:
            try:
                text = file_contents.decode(encoding)
                # Verify it's actually readable text (not binary garbage)
                if text and not self._is_binary_content(text):
                    return text
            except (UnicodeDecodeError, UnicodeError):
                continue

        raise ValueError(
            "Could not decode text file. "
            "Please ensure the file is saved with UTF-8 encoding."
        )

    def _is_binary_content(self, text: str) -> bool:
        """Check if decoded text is actually binary garbage."""
        if not text:
            return True
        # Count control characters (excluding common whitespace)
        control_chars = sum(1 for c in text[:1000] if ord(c) < 32 and c not in "\n\r\t")
        return control_chars > len(text[:1000]) * 0.1
