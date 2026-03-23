"""
Export Service
Handles document export to PDF/DOCX formats with Supabase
"""
from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime, timedelta
import io
import base64
import structlog

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph as RLParagraph, Spacer
from reportlab.lib.units import inch
from docx import Document as DocxDocument

from app.core.database import get_db, get_supabase, TABLES, SupabaseDB

logger = structlog.get_logger()


class ExportService:
    """Service for export operations using Supabase."""

    def __init__(self, db: Optional[SupabaseDB] = None):
        self.db = db or get_db()

    async def create_export(
        self,
        user_id: str,
        document_ids: List[str] = None,
        fmt: str = "pdf",
        filename: Optional[str] = None,
        options: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """Create an export bundle from application documents or standalone documents."""
        documents: List[Dict[str, Any]] = []

        # If options contains application_id, export from the applications table
        app_id = (options or {}).get("application_id")
        if app_id:
            app = await self.db.get(TABLES["applications"], app_id)
            if not app or app.get("user_id") != user_id:
                raise ValueError("Application not found or not accessible")

            doc_types = (options or {}).get("document_types", ["cv", "cover_letter"])
            type_map = {
                "cv": ("Tailored CV", "cv_html"),
                "cover_letter": ("Cover Letter", "cover_letter_html"),
                "personal_statement": ("Personal Statement", "personal_statement_html"),
                "portfolio": ("Portfolio", "portfolio_html"),
            }
            for dt in doc_types:
                if dt in type_map:
                    title, field = type_map[dt]
                    content = app.get(field, "")
                    if content:
                        documents.append({"title": title, "content": content, "format": "html"})
        elif document_ids:
            # Fallback: fetch from documents table
            for did in document_ids:
                doc = await self.db.get(TABLES["documents"], did)
                if not doc or doc.get("user_id") != user_id:
                    raise ValueError(f"Document {did} not found or not accessible")
                documents.append(doc)

        if not documents:
            raise ValueError("No documents to export")

        if not filename:
            timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
            filename = f"hirestack_export_{timestamp}.{fmt}"

        # Generate file content
        if fmt == "pdf":
            file_bytes = self._generate_pdf(documents, options)
        elif fmt == "docx":
            file_bytes = self._generate_docx(documents, options)
        elif fmt == "markdown":
            file_bytes = self._generate_markdown(documents)
        else:
            raise ValueError(f"Unsupported format: {fmt}")

        # Upload to Supabase Storage (preferred), fall back to base64 in DB
        file_url = await self._upload_to_storage(user_id, filename, file_bytes, fmt)

        record = {
            "user_id": user_id,
            "document_ids": document_ids or [],
            "format": fmt,
            "filename": filename,
            "file_size": len(file_bytes),
            "file_url": file_url,
            "options": options,
            "status": "completed",
        }
        doc_id = await self.db.create(TABLES["exports"], record)
        logger.info("export_created", export_id=doc_id, format=fmt, doc_count=len(documents))
        return await self.db.get(TABLES["exports"], doc_id)

    async def _upload_to_storage(self, user_id: str, filename: str, file_bytes: bytes, fmt: str) -> str:
        """Upload export file to Supabase Storage, returning a signed URL.
        Falls back to inline base64 data URI if Storage upload fails."""
        try:
            client = get_supabase()
            storage_path = f"{user_id}/exports/{filename}"
            content_types = {
                "pdf": "application/pdf",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "markdown": "text/markdown",
            }
            content_type = content_types.get(fmt, "application/octet-stream")

            # Upload to the 'uploads' bucket
            client.storage.from_("uploads").upload(
                path=storage_path,
                file=file_bytes,
                file_options={"content-type": content_type, "upsert": "true"},
            )

            # Generate a signed URL valid for 7 days
            signed = client.storage.from_("uploads").create_signed_url(
                storage_path, expires_in=7 * 24 * 3600
            )
            if signed and signed.get("signedURL"):
                logger.info("export_uploaded_to_storage", path=storage_path)
                return signed["signedURL"]

            # Fallback: public URL
            public = client.storage.from_("uploads").get_public_url(storage_path)
            return public
        except Exception as e:
            logger.warning("storage_upload_failed_falling_back_to_base64", error=str(e))
            # Fallback to base64 data URI
            b64 = base64.b64encode(file_bytes).decode()
            return f"data:application/octet-stream;base64,{b64}"

    def _strip_html(self, html: str) -> str:
        """Simple HTML to plain text conversion."""
        import re
        text = re.sub(r'<br\s*/?>', '\n', html)
        text = re.sub(r'</(p|div|h[1-6]|li|tr)>', '\n', text)
        text = re.sub(r'<[^>]+>', '', text)
        text = re.sub(r'&nbsp;', ' ', text)
        text = re.sub(r'&amp;', '&', text)
        text = re.sub(r'&lt;', '<', text)
        text = re.sub(r'&gt;', '>', text)
        text = re.sub(r'&#\d+;', '', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    # ── PDF ──
    def _generate_pdf(self, documents: List[Dict[str, Any]], options: Optional[Dict[str, Any]] = None) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle("CustomTitle", parent=styles["Heading1"], fontSize=18, spaceAfter=12)
        body_style = ParagraphStyle("CustomBody", parent=styles["Normal"], fontSize=11, spaceAfter=6, leading=14)
        story: list = []

        for document in documents:
            story.append(RLParagraph(document.get("title", "Untitled"), title_style))
            story.append(Spacer(1, 12))

            content = document.get("content", "")
            # If content is HTML, strip tags for PDF
            if "<" in content and ">" in content:
                content = self._strip_html(content)

            for para in content.split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                # Escape XML special chars for reportlab
                para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                story.append(RLParagraph(para, body_style))
            story.append(Spacer(1, 24))

        if not story:
            story.append(RLParagraph("No content to export.", body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()

    # ── DOCX ──
    def _generate_docx(self, documents: List[Dict[str, Any]], options: Optional[Dict[str, Any]] = None) -> bytes:
        docx = DocxDocument()
        for document in documents:
            docx.add_heading(document.get("title", "Untitled"), level=1)

            content = document.get("content", "")
            if "<" in content and ">" in content:
                content = self._strip_html(content)

            for para in content.split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                if para.startswith("#"):
                    level = min(len(para.split()[0].rstrip("#")), 4)
                    docx.add_heading(para.lstrip("#").strip(), level=level + 1)
                elif para.startswith("- ") or para.startswith("* "):
                    docx.add_paragraph(para[2:], style="List Bullet")
                else:
                    docx.add_paragraph(para)
            docx.add_page_break()

        buffer = io.BytesIO()
        docx.save(buffer)
        buffer.seek(0)
        return buffer.read()

    # ── Markdown ──
    def _generate_markdown(self, documents: List[Dict[str, Any]]) -> bytes:
        parts = []
        for document in documents:
            parts.append(f"# {document.get('title', 'Untitled')}\n\n")
            content = document.get("content", "")
            if "<" in content and ">" in content:
                content = self._strip_html(content)
            parts.append(content)
            parts.append("\n\n---\n\n")
        return "".join(parts).encode("utf-8")

    # ── CRUD ──
    async def get_user_exports(self, user_id: str, limit: int = 50) -> List[Dict[str, Any]]:
        return await self.db.query(
            TABLES["exports"],
            filters=[("user_id", "==", user_id)],
            order_by="created_at",
            order_direction="DESCENDING",
            limit=limit,
        )

    async def get_export(self, export_id: str, user_id: str) -> Optional[Dict[str, Any]]:
        export = await self.db.get(TABLES["exports"], export_id)
        if export and export.get("user_id") == user_id:
            return export
        return None

    async def download_export(self, export_id: str, user_id: str) -> Tuple[bytes, str, str]:
        """Return (file_bytes, filename, content_type)."""
        export = await self.get_export(export_id, user_id)
        if not export:
            raise ValueError("Export not found")

        file_url = export.get("file_url", "")
        content_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "markdown": "text/markdown",
        }
        content_type = content_types.get(export.get("format", ""), "application/octet-stream")
        filename = export.get("filename", "export")

        # Legacy base64 data URI
        if file_url.startswith("data:"):
            file_content = base64.b64decode(file_url.split(",")[1])
            return file_content, filename, content_type

        # Supabase Storage signed/public URL — fetch via httpx
        if file_url.startswith("http"):
            # SSRF protection: only allow URLs from our Supabase instance
            from urllib.parse import urlparse
            parsed = urlparse(file_url)
            allowed_hosts = ["dkfmcnfhvbqwsgpkgoag.supabase.co", "supabase.co"]
            if not parsed.hostname or not any(h in parsed.hostname for h in allowed_hosts):
                raise ValueError("Export file URL is not from a trusted source")
            import httpx
            async with httpx.AsyncClient() as client:
                resp = await client.get(file_url, follow_redirects=False)
                resp.raise_for_status()
                return resp.content, filename, content_type

        raise ValueError("Export file not available")

    async def delete_export(self, export_id: str, user_id: str) -> bool:
        export = await self.get_export(export_id, user_id)
        if not export:
            return False
        await self.db.delete(TABLES["exports"], export_id)
        return True
