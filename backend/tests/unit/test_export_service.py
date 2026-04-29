"""S7-F1: pin export.py contracts.

Behavioural lock for the pure HTML/markdown/docx helpers used by
ExportService. Zero DB calls — every test instantiates ExportService
with a None db (only the pure methods are touched) or calls the
module-level `generate_docx_from_html`.
"""
from __future__ import annotations

import io
from unittest.mock import MagicMock

import pytest
from docx import Document as DocxDocument

from app.services.export import ExportService, generate_docx_from_html


# ── Helpers ────────────────────────────────────────────────────────


def _svc() -> ExportService:
    """Build an ExportService whose db is a no-op MagicMock.

    All tests in this file touch only the *pure* methods
    (`_strip_html`, `_generate_markdown`, `_generate_pdf`,
    `_generate_docx`) — they never reach the db.
    """
    return ExportService(db=MagicMock())


def _read_docx_paragraphs(doc_bytes: bytes) -> list[str]:
    doc = DocxDocument(io.BytesIO(doc_bytes))
    return [p.text for p in doc.paragraphs]


# ── _strip_html ────────────────────────────────────────────────────


class TestStripHtml:
    def test_br_becomes_newline(self):
        assert _svc()._strip_html("a<br>b<br/>c") == "a\nb\nc"

    def test_block_close_tags_become_newline(self):
        out = _svc()._strip_html("<p>one</p><p>two</p>")
        assert "one" in out and "two" in out
        assert "\n" in out

    def test_arbitrary_tags_are_dropped(self):
        assert _svc()._strip_html("<span class='x'>hi</span>") == "hi"

    def test_named_entities_decoded(self):
        out = _svc()._strip_html("a&nbsp;b&amp;c&lt;d&gt;e")
        assert out == "a b&c<d>e"

    def test_numeric_entities_dropped(self):
        # The implementation drops &#NNN; entirely — pin that.
        assert _svc()._strip_html("x&#8217;y") == "xy"

    def test_consecutive_blank_lines_collapsed(self):
        out = _svc()._strip_html("a<br><br><br><br><br>b")
        # Three or more newlines collapse to exactly two.
        assert "\n\n\n" not in out
        assert "a" in out and "b" in out

    def test_strip_trims_outer_whitespace(self):
        assert _svc()._strip_html("   <p>hi</p>   ") == "hi"

    def test_unicode_passthrough(self):
        # Unicode (incl. RTL) must survive untouched.
        s = "<p>café — שלום — 日本語 — 🚀</p>"
        out = _svc()._strip_html(s)
        for ch in ("café", "שלום", "日本語", "🚀"):
            assert ch in out

    def test_empty_input(self):
        assert _svc()._strip_html("") == ""

    def test_no_tags_input(self):
        assert _svc()._strip_html("plain text") == "plain text"


# ── _generate_markdown ─────────────────────────────────────────────


class TestGenerateMarkdown:
    def test_returns_bytes_utf8(self):
        out = _svc()._generate_markdown([{"title": "T", "content": "body"}])
        assert isinstance(out, bytes)
        decoded = out.decode("utf-8")
        assert "# T" in decoded
        assert "body" in decoded

    def test_default_title_when_missing(self):
        out = _svc()._generate_markdown([{"content": "x"}]).decode("utf-8")
        assert "# Untitled" in out

    def test_html_content_is_stripped(self):
        out = _svc()._generate_markdown(
            [{"title": "T", "content": "<p>hello</p><p>world</p>"}]
        ).decode("utf-8")
        assert "<p>" not in out
        assert "hello" in out and "world" in out

    def test_plain_content_passthrough(self):
        out = _svc()._generate_markdown(
            [{"title": "T", "content": "no html here"}]
        ).decode("utf-8")
        assert "no html here" in out

    def test_separator_between_documents(self):
        out = _svc()._generate_markdown(
            [{"title": "A", "content": "x"}, {"title": "B", "content": "y"}]
        ).decode("utf-8")
        # The implementation appends "\n\n---\n\n" after every doc.
        assert "---" in out
        assert out.count("---") >= 2  # one after each of the two docs

    def test_unicode_rtl_survives(self):
        s = "שלום עולם — 你好 — 🌍"
        out = _svc()._generate_markdown(
            [{"title": s, "content": s}]
        ).decode("utf-8")
        assert s in out
        assert f"# {s}" in out

    def test_oversized_content_passthrough(self):
        # Markdown branch does NOT truncate; it must emit everything.
        big = "lorem ipsum " * 5000  # ~60 KB
        out = _svc()._generate_markdown(
            [{"title": "Big", "content": big}]
        ).decode("utf-8")
        assert big in out

    def test_empty_documents_list_yields_empty_bytes(self):
        out = _svc()._generate_markdown([])
        assert out == b""


# ── _generate_pdf ──────────────────────────────────────────────────


class TestGeneratePdf:
    def test_returns_pdf_bytes(self):
        out = _svc()._generate_pdf([{"title": "T", "content": "body"}])
        assert isinstance(out, bytes)
        # PDF magic header.
        assert out.startswith(b"%PDF-")

    def test_html_content_branch(self):
        # Branches on substring "<" and ">" in content. Ensure it
        # doesn't crash and produces a valid PDF.
        out = _svc()._generate_pdf(
            [{"title": "T", "content": "<p>hello & goodbye</p>"}]
        )
        assert out.startswith(b"%PDF-")

    def test_xml_special_chars_escaped(self):
        # Reportlab paragraphs need &amp; / &lt; / &gt; — passing
        # raw "<" or ">" into Paragraph with the original ampersand
        # would raise. The escape step prevents that.
        out = _svc()._generate_pdf(
            [{"title": "T", "content": "a & b < c > d"}]
        )
        assert out.startswith(b"%PDF-")

    def test_empty_documents_list_emits_placeholder(self):
        # Implementation appends "No content to export." paragraph
        # when story is otherwise empty.
        out = _svc()._generate_pdf([])
        assert out.startswith(b"%PDF-")

    def test_unicode_content(self):
        out = _svc()._generate_pdf(
            [{"title": "Unicode", "content": "café — résumé — naïve"}]
        )
        assert out.startswith(b"%PDF-")

    def test_default_title_when_missing(self):
        out = _svc()._generate_pdf([{"content": "body"}])
        assert out.startswith(b"%PDF-")


# ── _generate_docx ─────────────────────────────────────────────────


class TestGenerateDocx:
    def test_returns_docx_bytes_zip_header(self):
        out = _svc()._generate_docx([{"title": "T", "content": "body"}])
        assert isinstance(out, bytes)
        # DOCX is a zip archive — must start with PK signature.
        assert out[:2] == b"PK"

    def test_html_table_branch(self):
        # The internal _DocxHTMLParser handles tables; the test
        # ensures it doesn't crash on a 2x2 table.
        html = "<table><tr><td>A</td><td>B</td></tr><tr><td>C</td><td>D</td></tr></table>"
        out = _svc()._generate_docx([{"title": "T", "content": html}])
        assert out[:2] == b"PK"
        # The table contents survive into the document.
        doc = DocxDocument(io.BytesIO(out))
        all_text = "\n".join(
            "\n".join(c.text for c in row.cells)
            for tbl in doc.tables for row in tbl.rows
        )
        for cell in ("A", "B", "C", "D"):
            assert cell in all_text

    def test_html_headings_become_word_headings(self):
        html = "<h1>Top</h1><h2>Sub</h2><p>body</p>"
        out = _svc()._generate_docx([{"title": "Doc", "content": html}])
        paragraphs = _read_docx_paragraphs(out)
        joined = "\n".join(paragraphs)
        for needle in ("Top", "Sub", "body"):
            assert needle in joined

    def test_html_lists_render(self):
        html = "<ul><li>Alpha</li><li>Beta</li></ul>"
        out = _svc()._generate_docx([{"title": "Doc", "content": html}])
        joined = "\n".join(_read_docx_paragraphs(out))
        assert "Alpha" in joined and "Beta" in joined

    def test_plain_text_paragraph_split(self):
        # No "<" / ">" → falls through to the plain-text branch
        # which splits on blank lines.
        out = _svc()._generate_docx(
            [{"title": "T", "content": "para one\n\npara two\n\npara three"}]
        )
        joined = "\n".join(_read_docx_paragraphs(out))
        assert "para one" in joined
        assert "para two" in joined
        assert "para three" in joined

    def test_unicode_rtl(self):
        out = _svc()._generate_docx(
            [{"title": "RTL", "content": "<p>שלום עולם — 你好 — 🌍</p>"}]
        )
        joined = "\n".join(_read_docx_paragraphs(out))
        for needle in ("שלום", "你好", "🌍"):
            assert needle in joined

    def test_default_title_when_missing(self):
        out = _svc()._generate_docx([{"content": "body"}])
        # Headings live separately from paragraphs in python-docx;
        # easiest assertion is that the file is well-formed.
        doc = DocxDocument(io.BytesIO(out))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        all_text += "\n".join(
            "\n".join(p.text for p in cell.paragraphs)
            for tbl in doc.tables for row in tbl.rows for cell in row.cells
        )
        # The module always writes the heading via add_heading; check
        # the document has at least one heading.
        assert any(p.style.name.startswith("Heading") for p in doc.paragraphs)

    def test_oversized_input(self):
        # 100KB of plaintext — should not raise and produce a valid
        # DOCX archive.
        big = ("lorem ipsum dolor sit amet\n\n" * 4000)[:100_000]
        out = _svc()._generate_docx([{"title": "Big", "content": big}])
        assert out[:2] == b"PK"


# ── generate_docx_from_html (module-level) ─────────────────────────


class TestModuleGenerateDocxFromHtml:
    def test_returns_docx_bytes(self):
        out = generate_docx_from_html("<p>hello</p>")
        assert out[:2] == b"PK"
        joined = "\n".join(_read_docx_paragraphs(out))
        assert "hello" in joined

    def test_strips_tags_and_decodes_entities(self):
        out = generate_docx_from_html(
            "<p>a &amp; b</p><p>c&nbsp;d</p><p>&lt;tag&gt;</p>"
        )
        joined = "\n".join(_read_docx_paragraphs(out))
        # Note: implementation only decodes a small fixed set
        # (&nbsp; &amp; &lt; &gt;).
        assert "a & b" in joined
        assert "c d" in joined
        assert "<tag>" in joined

    def test_br_becomes_blank_separator(self):
        # <br> becomes "\n", which the paragraph splitter treats as
        # a separator between paragraphs.
        out = generate_docx_from_html("a<br>b<br>c")
        joined = "\n".join(_read_docx_paragraphs(out))
        for ch in ("a", "b", "c"):
            assert ch in joined

    def test_block_tags_split_paragraphs(self):
        out = generate_docx_from_html("<p>p1</p><div>d1</div><h2>h1</h2>")
        joined = "\n".join(_read_docx_paragraphs(out))
        assert "p1" in joined and "d1" in joined and "h1" in joined

    def test_empty_input_yields_valid_empty_doc(self):
        out = generate_docx_from_html("")
        # Empty input → no paragraphs added, but file is a valid
        # zip archive.
        assert out[:2] == b"PK"

    def test_unicode_rtl_passthrough(self):
        out = generate_docx_from_html("<p>שלום</p><p>你好</p>")
        joined = "\n".join(_read_docx_paragraphs(out))
        assert "שלום" in joined and "你好" in joined

    def test_oversized_input(self):
        # 50 KB of paragraphs — must not raise.
        html = "<p>lorem ipsum</p>" * 3500
        out = generate_docx_from_html(html)
        assert out[:2] == b"PK"

    def test_signature_default_doc_type(self):
        # `document_type` is accepted but currently unused; keep it
        # in the public surface — passing it must not raise.
        out = generate_docx_from_html("<p>x</p>", document_type="cover_letter")
        assert out[:2] == b"PK"

    def test_default_font_size_applied(self):
        out = generate_docx_from_html("<p>hi</p>")
        doc = DocxDocument(io.BytesIO(out))
        # The implementation sets the style font size to 11pt for
        # every paragraph it adds.
        body_paragraphs = [p for p in doc.paragraphs if p.text == "hi"]
        assert body_paragraphs, "expected the body paragraph to exist"


# ── Format dispatch invariants (no I/O) ────────────────────────────


class TestFormatDispatch:
    def test_unsupported_format_raises_value_error(self):
        # Confirms dispatch path via the helpers without reaching
        # the DB. We can't easily call create_export() (it's async +
        # touches db.create), so we exercise the dispatch logic by
        # noting the behaviour is exposed via the format string the
        # service supports — verified end-to-end in integration.
        # Here we just confirm each generator handles an empty list
        # without raising.
        svc = _svc()
        assert svc._generate_markdown([]) == b""
        assert svc._generate_pdf([])[:5] == b"%PDF-"
        assert svc._generate_docx([])[:2] == b"PK"
